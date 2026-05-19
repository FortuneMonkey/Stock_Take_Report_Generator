import streamlit as st
import openpyxl
import io, os, json, re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.enums import TA_RIGHT, TA_CENTER

st.set_page_config(page_title="Stock Take Report Generator", page_icon="📦", layout="wide")
st.markdown("""
<style>
.main-header{background:linear-gradient(90deg,#1F4E79,#2E75B6);color:white;padding:1.4rem 2rem;border-radius:10px;margin-bottom:1.2rem}
.warn{background:#fff3cd;border:1px solid #ffeeba;padding:.7rem 1rem;border-radius:6px;color:#856404;margin-bottom:.8rem}
.ok{background:#d4edda;border:1px solid #c3e6cb;padding:.7rem 1rem;border-radius:6px;color:#155724;margin-bottom:.5rem}
.store-ok{color:#155724;font-weight:600}
.store-err{color:#721c24;font-weight:600}
</style>
""", unsafe_allow_html=True)
st.markdown('<div class="main-header"><h2 style="margin:0">📦 Stock Take Report Generator</h2>'
            '<p style="margin:0;opacity:.85">APRIL Fiber R&D — Quarterly Stock Take Automation</p></div>',
            unsafe_allow_html=True)

STORE_DEFS = [
    {"key":"SOL",  "name":"Analytical Lab. Store",  "default_pic":"Nanda Aprilinati", "sloc":"4413","seed":False},
    {"key":"BIO",  "name":"Biomol Lab. Store",      "default_pic":"Jonawi Ng",         "sloc":"4413","seed":False},
    {"key":"SEED", "name":"GD Seed Store",          "default_pic":"Ela",               "sloc":"4413","seed":True},
    {"key":"TCL",  "name":"KTC Store",              "default_pic":"Anding Oktaviani",  "sloc":"4414","seed":False},
    {"key":"PHL",  "name":"Plant Health Lab. Store","default_pic":"Deviana Tan",      "sloc":"4413","seed":False},
    {"key":"RDS",  "name":"R&D Main Store",         "default_pic":"Deni Sried",        "sloc":"4413","seed":False},
    {"key":"KRN",  "name":"R&D Nursery Store",      "default_pic":"Pebriansyah Bakti", "sloc":"4413","seed":False},

]
NM_DEFS = [{"sloc":"4413","desc":"R&D Stores"},{"sloc":"4414","desc":"KTC Store"}]
CONFIG_FILE = "config.json"

def fmt_idr(val):
    try: v=int(val); return "-" if v==0 else f"{v:,}".replace(",",".")
    except: return "-"

def safe_int(v):
    if v is None: return 0
    if isinstance(v,(int,float)): return int(v)
    s=str(v).strip(); neg=s.startswith("-")
    s=s.lstrip("-").replace("Rp","").replace("rp","").replace(".","").replace(",","").strip()
    try: return -int(s) if neg else int(s)
    except: return 0

def safe_float(v):
    if v is None: return 0.0
    if isinstance(v,(int,float)): return float(v)
    s=str(v).strip(); neg=s.startswith("-")
    s=s.lstrip("-").replace("Rp","").replace(",","").strip()
    try: return -float(s) if neg else float(s)
    except: return 0.0

def safe_str(v): return str(v).strip() if v is not None else ""

def load_config():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE) as f: return json.load(f)
    return {"quarter":"Q3","year":"2025","start_date":"22-Sep-2025","end_date":"25-Sep-2025",
            "prepared_by":"Susanna Chitraresmi","acknowledged_by":"Sabar T. H. Siregar",
            "approved_by":"Alvaro J. Duran S.","approved_title":"APRIL Fiber R&D Head",
            "non_moving":[{"sloc":"4413","desc":"R&D Stores","count":72,"value":283043745},
                          {"sloc":"4414","desc":"KTC Store","count":92,"value":470126900}],
            "store_overrides":{}}

def save_config(cfg):
    with open(CONFIG_FILE,"w") as f: json.dump(cfg,f,indent=2)

def extract_quarter_from_filename(fn):
    m=re.search(r'Q(\d)',fn.upper()); return f"Q{m.group(1)}" if m else None

def extract_year_from_filename(fn):
    m=re.search(r'(20\d{2})',fn); return m.group(1) if m else None

def extract_date_from_filename(fn):
    """Extract stock take date e.g. 'STOCK TAKE BIO 22-Sep-2025 Q4.xlsx' -> '22-Sep-2025'"""
    m=re.search(r'(\d{1,2}-[A-Za-z]{3}-20\d{2})', fn)
    return m.group(1) if m else None

def detect_store(filename):
    fn=filename.upper()
    for sd in STORE_DEFS:
        if sd["key"] in fn: return sd
    return None

def detect_nm_file(filename):
    fn=filename.upper()
    if not any(k in fn for k in ["NON","MOVING","NMS"]): return None
    for nm in NM_DEFS:
        if nm["sloc"] in fn: return nm
    return None

def read_nm_auto(ws):
    header_found=False; rows=[]; total_mc=0; total_value=0.0; empty_streak=0
    for row in ws.iter_rows(min_row=1,max_row=20,values_only=True):
        if row is None or all(c is None for c in row):
            empty_streak+=1
            if empty_streak>3: break
            continue
        empty_streak=0
        row_strs=[str(c).strip().lower() if c is not None else "" for c in row]
        if any(h in row_strs for h in ["number of mc","need","criteria","no"]):
            header_found=True; continue
        if not header_found: continue
        criteria=str(row[1]).strip() if row[1] is not None else ""
        if not criteria: continue
        is_total=criteria.lower()=="total"
        mc_need=safe_int(row[2]) if len(row)>2 else 0
        mc_dont=safe_int(row[3]) if len(row)>3 else 0
        val_need=safe_float(row[4]) if len(row)>4 else 0.0
        val_dont=safe_float(row[5]) if len(row)>5 else 0.0
        if is_total:
            total_mc=mc_need+mc_dont; total_value=val_need+val_dont; break
        else:
            rows.append({"criteria":criteria,"mc_need":mc_need,"mc_dont":mc_dont,"mc_total":mc_need+mc_dont,
                         "val_need":val_need,"val_dont":val_dont,"val_total":val_need+val_dont})
    if total_mc==0 and rows:
        total_mc=sum(r["mc_total"] for r in rows); total_value=sum(r["val_total"] for r in rows)
    return {"count":total_mc,"value":int(total_value),"rows":rows}

def parse_nm_file(uploaded_file, nm_def):
    wb=openpyxl.load_workbook(uploaded_file,data_only=True,read_only=True)
    auto_ws=None
    for name in wb.sheetnames:
        if name.strip().lower()=="auto": auto_ws=wb[name]; break
    if auto_ws is None:
        wb.close(); return None,f"Sheet 'auto' not found. Available: {', '.join(wb.sheetnames)}"
    data=read_nm_auto(auto_ws); wb.close()
    data["sloc"]=nm_def["sloc"]; data["desc"]=nm_def["desc"]
    return data,None

def read_standard_auto(ws):
    result={}
    label_map={"balance":["balance","balance items"],"surplus":["surplus","surplus items"],
               "deficit":["deficit","deficit items"],"total":["total"]}
    for row in ws.iter_rows(min_row=1,values_only=True):
        if row is None: continue
        for col_idx,cell in enumerate(row):
            if cell is None: continue
            cell_str=str(cell).strip().lower()
            for key,patterns in label_map.items():
                if any(cell_str==p for p in patterns):
                    nums=[int(v) for v in row if isinstance(v,(int,float)) and v is not None]
                    if len(nums)>=3:
                        if nums[0] in [1,2,3] and key!="total": nums=nums[1:]
                        result[key]={"mc":nums[0] if len(nums)>0 else 0,"sap":nums[1] if len(nums)>1 else 0,
                                     "actual":nums[2] if len(nums)>2 else 0,"var":nums[3] if len(nums)>3 else 0}
                    break
    return result

def read_seed_auto(ws):
    header_row_idx=None; code_col=sap_val_col=actual_val_col=var_idr_col=None
    for r_idx,row in enumerate(ws.iter_rows(values_only=True),start=1):
        if row is None: continue
        rs=[str(c).strip().lower() if c is not None else "" for c in row]
        if "code" in rs and any("sap" in s for s in rs):
            header_row_idx=r_idx
            for i,h in enumerate(rs):
                if h=="code": code_col=i
                if ("sap value" in h or h=="sap value (idr)") and "actual" not in h: sap_val_col=i
                if "actual value" in h: actual_val_col=i
                if "var" in h and "idr" in h: var_idr_col=i
            break
    if header_row_idx is None: header_row_idx=1
    if code_col       is None: code_col=10
    if sap_val_col    is None: sap_val_col=7
    if actual_val_col is None: actual_val_col=8
    if var_idr_col    is None: var_idr_col=9
    groups={"balance":{"mc":0,"sap":0,"actual":0,"var":0},"surplus":{"mc":0,"sap":0,"actual":0,"var":0},
            "deficit":{"mc":0,"sap":0,"actual":0,"var":0}}
    species_rows=[]; total_from_sheet=None; row_num=0
    no_col=desc_col=unit_col=sap_qty_col=act_qty_col=var_qty_col=mc_col=None
    for r_idx,row in enumerate(ws.iter_rows(values_only=True),start=1):
        if r_idx==header_row_idx:
            rs=[str(c).strip().lower() if c is not None else "" for c in row]
            for i,h in enumerate(rs):
                if h in ("no","no.","#"): no_col=i
                if h=="mc": mc_col=i
                if "material description" in h or h=="description": desc_col=i
                if h=="unit": unit_col=i
                if "sap qty" in h and "value" not in h: sap_qty_col=i
                if ("actual qty" in h or "total actual qty" in h) and "value" not in h: act_qty_col=i
                if "var" in h and "idr" not in h and i<(sap_val_col or 99): var_qty_col=i
            break
    if sap_qty_col  is None: sap_qty_col=4
    if act_qty_col  is None: act_qty_col=5
    if var_qty_col  is None: var_qty_col=6
    if no_col       is None: no_col=0
    if mc_col       is None: mc_col=1
    if desc_col     is None: desc_col=2
    if unit_col     is None: unit_col=3
    for row in ws.iter_rows(min_row=header_row_idx+1,values_only=True):
        if row is None or all(c is None for c in row): continue
        first_text=next((str(c).strip().lower() for c in row if c is not None and
                         not str(c).strip().lstrip("-").replace("Rp","").replace(".","").replace(",","").isdigit()),"")
        if first_text=="total":
            total_from_sheet={"mc":0,
                "sap":safe_int(row[sap_val_col]) if sap_val_col is not None and len(row)>sap_val_col else 0,
                "actual":safe_int(row[actual_val_col]) if actual_val_col is not None and len(row)>actual_val_col else 0,
                "var":safe_int(row[var_idr_col]) if var_idr_col is not None and len(row)>var_idr_col else 0,
                "sap_qty":safe_int(row[sap_qty_col]) if sap_qty_col and len(row)>sap_qty_col else 0,
                "actual_qty":safe_int(row[act_qty_col]) if act_qty_col and len(row)>act_qty_col else 0,
                "var_qty":safe_int(row[var_qty_col]) if var_qty_col and len(row)>var_qty_col else 0}
            continue
        if code_col is None or len(row)<=code_col: continue
        code_raw=safe_str(row[code_col]).lower()
        if not code_raw: continue
        if   "balance" in code_raw: grp="balance"
        elif "surplus" in code_raw: grp="surplus"
        elif "deficit" in code_raw: grp="deficit"
        else: continue
        row_num+=1
        sap_v=safe_int(row[sap_val_col]) if sap_val_col is not None and len(row)>sap_val_col else 0
        actual_v=safe_int(row[actual_val_col]) if actual_val_col is not None and len(row)>actual_val_col else 0
        var_v=safe_int(row[var_idr_col]) if var_idr_col is not None and len(row)>var_idr_col else 0
        sap_qty_v=safe_int(row[sap_qty_col]) if sap_qty_col is not None and len(row)>sap_qty_col else 0
        actual_qty_v=safe_int(row[act_qty_col]) if act_qty_col is not None and len(row)>act_qty_col else 0
        var_qty_v=safe_int(row[var_qty_col]) if var_qty_col is not None and len(row)>var_qty_col else 0
        species_rows.append({"no":row_num,"mc":safe_str(row[mc_col]) if mc_col is not None and len(row)>mc_col else "",
                             "desc":safe_str(row[desc_col]) if desc_col is not None and len(row)>desc_col else "",
                             "unit":safe_str(row[unit_col]) if unit_col is not None and len(row)>unit_col else "",
                             "sap_qty":sap_qty_v,"actual_qty":actual_qty_v,"var_qty":var_qty_v,
                             "sap":sap_v,"actual":actual_v,"var":var_v,"code":grp})
        groups[grp]["mc"]+=1; groups[grp]["sap"]+=sap_v; groups[grp]["actual"]+=actual_v; groups[grp]["var"]+=var_v
    mc_total=sum(groups[k]["mc"] for k in groups)
    if total_from_sheet is not None:
        total_from_sheet["mc"]=mc_total; total=total_from_sheet
    else:
        total={"mc":mc_total,"sap":sum(groups[k]["sap"] for k in groups),
               "actual":sum(groups[k]["actual"] for k in groups),"var":sum(groups[k]["var"] for k in groups),
               "sap_qty":sum(r["sap_qty"] for r in species_rows),
               "actual_qty":sum(r["actual_qty"] for r in species_rows),
               "var_qty":sum(r["var_qty"] for r in species_rows)}
    return {**groups,"total":total,"species_rows":species_rows}

def parse_store_file(uploaded_file, store_def):
    wb=openpyxl.load_workbook(uploaded_file,data_only=True)
    auto_ws=None
    for name in wb.sheetnames:
        if name.strip().lower()=="auto": auto_ws=wb[name]; break
    if auto_ws is None:
        return None,f"Sheet 'auto' not found. Available: {', '.join(wb.sheetnames)}"
    if store_def["seed"]:
        data=read_seed_auto(auto_ws)
        seed_lots=0
        for name in wb.sheetnames:
            if "material master" in name.strip().lower():
                for row in wb[name].iter_rows(min_row=2,values_only=True):
                    if any(c is not None for c in row): seed_lots+=1
                break
        data["seed_lots_count"]=seed_lots
        data["species_count"]=len(data.get("species_rows",[]))
        total_g=sum(r["sap_qty"] for r in data.get("species_rows",[]))
        data["total_weight_tons"]=f"{total_g/1_000_000:.1f} tons"
        data["total_sap_qty_g"]=total_g
    else:
        data=read_standard_auto(auto_ws)
    if not data: return None,"Could not read data from 'auto' sheet"
    q=extract_quarter_from_filename(uploaded_file.name)
    if q: data["quarter_from_file"]=q
    y=extract_year_from_filename(uploaded_file.name)
    if y: data["year_from_file"]=y
    d=extract_date_from_filename(uploaded_file.name)
    if d: data["date_from_file"]=d
    return data,None

def auto_remarks(store_name, data, cfg):
    total=data.get("total",{}); total_var=total.get("var",0)
    quarter=cfg.get("quarter","Q3")
    overrides=cfg.get("store_overrides",{}).get(store_name,{})
    perf_ov=overrides.get("performance_note",None); act_ov=overrides.get("action_note",None)
    if total_var==0:
        perf=perf_ov if perf_ov is not None else f"{store_name} successfully managed to perform Zero Variance at {quarter} stock takes."
        action=act_ov if act_ov is not None else ""
    elif total_var>0:
        perf=perf_ov if perf_ov is not None else ""
        action=act_ov if act_ov is not None else "Must pay more attention on handling the issued material to requestor and make sure all material is handed over."
    else:
        perf=perf_ov if perf_ov is not None else ""
        action=act_ov if act_ov is not None else "Must to be strict on the administration of the material movement."
    return perf,action

def build_intro(store_name, data, store_def, cfg):
    total=data.get("total",{}); mc=total.get("mc",0); sap=total.get("sap",0)
    actual=total.get("actual",0); var=total.get("var",0)
    if store_def["seed"]:
        lots=str(data.get("seed_lots_count","")); species=str(data.get("species_count","")); weight=data.get("total_weight_tons","")
        if var==0:
            vl=(f"Total value of the seed stock on SAP worth IDR {fmt_idr(sap)},- where the actual value was IDR {fmt_idr(actual)},- and the quantity and value are matched.")
        elif var>0:
            vl=(f"Total value of the seed stock on SAP worth IDR {fmt_idr(sap)},- where the actual value was IDR {fmt_idr(actual)},- and surplus value was IDR {fmt_idr(abs(var))},-.")
        else:
            vl=(f"Total value of the seed stock on SAP worth IDR {fmt_idr(sap)},- where the actual value was IDR {fmt_idr(actual)},- and deficit value was IDR {fmt_idr(abs(var))},-.")
        return (f"Current number of Seed Lots maintained in {store_name} are {lots} seed lots from {species} Species with total weight \u00b1 {weight}. "
                f"These seed lots recorded on SAP by species based.\n\n{vl}")
    else:
        if var==0:   vt="where the actual quantity and value are matched."
        elif var>0:  vt=f"where the actual quantity was IDR {fmt_idr(actual)} and total surplus IDR {fmt_idr(var)}."
        else:        vt=f"where the actual quantity IDR {fmt_idr(actual)} and total deficit was IDR {fmt_idr(abs(var))},-"
        return f"Current number of materials kept in {store_name} are {mc} items with SAP total value worth IDR {fmt_idr(sap)},- {vt}"

def set_cell_bg(cell,hex_color):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); shd=OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hex_color); tcPr.append(shd)

def set_cell_borders(cell,color="CCCCCC"):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcBorders=OxmlElement("w:tcBorders")
    for side in ["top","left","bottom","right"]:
        b=OxmlElement(f"w:{side}"); b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"4")
        b.set(qn("w:space"),"0"); b.set(qn("w:color"),color); tcBorders.append(b)
    tcPr.append(tcBorders)

def no_border(cell):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcBorders=OxmlElement("w:tcBorders")
    for side in ["top","left","bottom","right"]:
        b=OxmlElement(f"w:{side}"); b.set(qn("w:val"),"none"); tcBorders.append(b)
    tcPr.append(tcBorders)

def add_run(para,text,bold=False,size=10,color=None):
    r=para.add_run(text); r.font.bold=bold; r.font.size=Pt(size)
    if color: r.font.color.rgb=RGBColor(*color)
    return r

def _add_header_footer(doc):
    for section in doc.sections:
        header=section.header; header.is_linked_to_previous=False
        hp=header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear(); hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        hr=hp.add_run("Internal"); hr.font.bold=True; hr.font.size=Pt(10); hr.font.name="Calibri"
        footer=section.footer; footer.is_linked_to_previous=False
        fp=footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=fp.add_run(); run.font.size=Pt(9)
        fldChar1=OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"),"begin")
        instrText=OxmlElement("w:instrText"); instrText.text="PAGE"; instrText.set(qn("xml:space"),"preserve")
        fldChar2=OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"),"end")
        run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
        run2=fp.add_run(" of "); run2.font.size=Pt(9)
        run3=fp.add_run(); run3.font.size=Pt(9)
        fldChar3=OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"),"begin")
        instrText2=OxmlElement("w:instrText"); instrText2.text="NUMPAGES"; instrText2.set(qn("xml:space"),"preserve")
        fldChar4=OxmlElement("w:fldChar"); fldChar4.set(qn("w:fldCharType"),"end")
        run3._r.append(fldChar3); run3._r.append(instrText2); run3._r.append(fldChar4)

def generate_docx(cfg, stores_results):
    doc=Document()
    for section in doc.sections:
        section.top_margin=Cm(3.0); section.bottom_margin=Cm(3.0)
        section.left_margin=Cm(3.0); section.right_margin=Cm(2.5)
    doc.styles["Normal"].font.name="Calibri"; doc.styles["Normal"].font.size=Pt(10)
    _add_header_footer(doc)
    quarter=cfg["quarter"]; year=cfg["year"]
    report_date=datetime.now().strftime("%d %B %Y").lstrip("0")
    start_date=cfg["start_date"]; end_date=cfg["end_date"]
    approved_by=cfg["approved_by"]; approved_ttl=cfg["approved_title"]
    prepared_by=cfg["prepared_by"]; acknowledged=cfg["acknowledged_by"]

    def para(text="",bold=False,size=10,align=WD_ALIGN_PARAGRAPH.LEFT,before=0,after=4):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
        p.alignment=align
        if text: add_run(p,text,bold=bold,size=size)
        return p

    para(f"Pangkalan Kerinci, {report_date}",before=4)
    p=para(before=8,after=3); add_run(p,"Subject: ",bold=True); add_run(p,"Material Stock Take Report")
    p=para(before=3,after=3); add_run(p,"To: ",bold=True); add_run(p,f"Mr. {approved_by}")
    para(approved_ttl,before=0,after=10)
    para("Dear Sir,",before=4,after=6)
    para(f"Material Stock Take {quarter} at all R&D Stores had been done from {start_date} to {end_date}. "
         "The activity from Analytic Lab. Store, Biomolecular Lab., GD Seed Store, "
         "KTC Store., Plant Health Lab. Store, R&D Main Store, and R&D Nursery Store.",before=2,after=10)

    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(4)
    add_run(p,"A.\t",bold=True); add_run(p,"Code and Termination",bold=True)
    for level,runs in [
        (0,[("Store Location on SAP, R&D has two store locations:",False)]),
        (1,[("4413\t",False),("maintain all data from Genetic Deployment Seed Store, Analytic Lab. Store, Plant Health Lab. Store, Biomolecular Lab. R&D Main Store, and R&D Nursery Store.",False)]),
        (1,[("4414\t",False),("maintain data from RGE Tissue Culture Lab. and KTC Store.",False)]),
        (0,[("Balance",True),(" is where Actual Stock matched with SAP stock.",False)]),
        (0,[("Surplus",True),(" is where Actual Stock exceed from SAP stock.",False)]),
        (0,[("Deficit",True),(" is where Actual Stock less from SAP stock.",False)]),
        (0,[("Stock Take Procedure",False)]),
        (1,[("Basis data used is from SAP.",False)]),
        (1,[("Material that has been using before Stock Take with proven manual SR will be recorded SR process.",False)]),
        (1,[("Material which failed to be shown is recorded as none.",False)]),
    ]:
        p=doc.add_paragraph(style="List Bullet" if level==0 else "List Bullet 2")
        p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
        for text,bold in runs: add_run(p,text,bold=bold)

    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4)
    add_run(p,"B.\t",bold=True); add_run(p,"Stock Take Summary",bold=True)
    col_widths=[Cm(1.0),Cm(4.5),Cm(2.5),Cm(3.5),Cm(3.5),Cm(3.0)]
    aligns=[WD_ALIGN_PARAGRAPH.CENTER,WD_ALIGN_PARAGRAPH.LEFT,WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT,WD_ALIGN_PARAGRAPH.RIGHT,WD_ALIGN_PARAGRAPH.RIGHT]

    for idx,store_def in enumerate(STORE_DEFS,start=1):
        sname=store_def["name"]; result=stores_results.get(store_def["key"])
        if result is None: continue
        data=result["data"]; overrides=cfg.get("store_overrides",{}).get(sname,{})
        pic=overrides.get("pic",store_def["default_pic"])
        date=overrides.get("stock_take_date") or data.get("date_from_file") or cfg["start_date"]
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(3)
        add_run(p,f"{idx}.\t{sname}, {date}, PIC {pic}.")
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(6)
        add_run(p,build_intro(sname,data,store_def,cfg))
        total=data.get("total",{}); is_seed=store_def["seed"]
        if is_seed:
            sr=data.get("species_rows",[])
            sh=["No.","MC","Material Description","Unit","SAP Qty","Total Actual\nQty. (g)","Var.\n(g)","SAP Value\n(IDR)","Actual Value\n(IDR)","Var. (IDR)","Code"]
            sw=[Cm(0.7),Cm(1.4),Cm(3.8),Cm(1.0),Cm(1.8),Cm(1.8),Cm(1.3),Cm(2.6),Cm(2.6),Cm(1.8),Cm(1.5)]
            sa=[WD_ALIGN_PARAGRAPH.CENTER]*4+[WD_ALIGN_PARAGRAPH.RIGHT]*7
            tbl=doc.add_table(rows=1+len(sr)+1,cols=11); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl.style="Table Grid"
            for i,(cell,text,w) in enumerate(zip(tbl.rows[0].cells,sh,sw)):
                cell.width=w; p2=cell.paragraphs[0]; p2.clear()
                r=p2.add_run(text); r.font.bold=True; r.font.size=Pt(7.5); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
                p2.alignment=WD_ALIGN_PARAGRAPH.CENTER; cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
                set_cell_bg(cell,"2E75B6"); set_cell_borders(cell,"FFFFFF")
            for r_i,s in enumerate(sr,start=1):
                vq=s["var_qty"]; vv=s["var"]
                vals=[str(s["no"]),str(s["mc"]),s["desc"],s["unit"],fmt_idr(s["sap_qty"]),fmt_idr(s["actual_qty"]),
                      "-" if vq==0 else fmt_idr(vq),fmt_idr(s["sap"]),fmt_idr(s["actual"]),"-" if vv==0 else fmt_idr(vv),s["code"].capitalize()]
                for i,(cell,val,align,w) in enumerate(zip(tbl.rows[r_i].cells,vals,sa,sw)):
                    cell.width=w; p2=cell.paragraphs[0]; p2.clear(); rn=p2.add_run(val); rn.font.size=Pt(7.5)
                    if i==9 and vv<0: rn.font.color.rgb=RGBColor(0xC0,0x00,0x00)
                    p2.alignment=align; cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER; set_cell_borders(cell,"CCCCCC")
            tot=total; tvq=tot.get("var_qty",0); tv=tot.get("var",0)
            tv_vals=["","Total","","",fmt_idr(tot.get("sap_qty",0)),fmt_idr(tot.get("actual_qty",0)),
                     "-" if tvq==0 else fmt_idr(tvq),fmt_idr(tot.get("sap",0)),fmt_idr(tot.get("actual",0)),
                     "-" if tv==0 else fmt_idr(tv),""]
            for i,(cell,val,align,w) in enumerate(zip(tbl.rows[-1].cells,tv_vals,sa,sw)):
                cell.width=w; p2=cell.paragraphs[0]; p2.clear(); rn=p2.add_run(val); rn.font.bold=True; rn.font.size=Pt(7.5)
                if i==9 and tv<0: rn.font.color.rgb=RGBColor(0xC0,0x00,0x00)
                p2.alignment=align; cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
                set_cell_bg(cell,"D6E4F0"); set_cell_borders(cell,"AAAAAA")
        else:
            rs=[("Balance Items",data.get("balance",{})),("Surplus Items",data.get("surplus",{})),("Deficit Items",data.get("deficit",{}))]
            tbl=doc.add_table(rows=5,cols=6); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl.style="Table Grid"
            hdr=["No.","Description","Number\nof MC","SAP Value\n(IDR)","Actual Value\n(IDR)","Var. (IDR)"]
            for i,(cell,text,w) in enumerate(zip(tbl.rows[0].cells,hdr,col_widths)):
                cell.width=w; p2=cell.paragraphs[0]; p2.clear()
                r=p2.add_run(text); r.font.bold=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
                p2.alignment=WD_ALIGN_PARAGRAPH.CENTER; cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
                set_cell_bg(cell,"2E75B6"); set_cell_borders(cell,"FFFFFF")
            for r_i,(desc,rd) in enumerate(rs,start=1):
                vv=rd.get("var",0); mc_d=fmt_idr(rd.get("mc",0)) if rd.get("mc",0)!=0 else "-"
                vals=[str(r_i),desc,mc_d,fmt_idr(rd.get("sap",0)),fmt_idr(rd.get("actual",0)),fmt_idr(vv)]
                for i,(cell,val,align,w) in enumerate(zip(tbl.rows[r_i].cells,vals,aligns,col_widths)):
                    cell.width=w; p2=cell.paragraphs[0]; p2.clear(); rn=p2.add_run(val); rn.font.size=Pt(9)
                    if i==5 and vv<0: rn.font.color.rgb=RGBColor(0xC0,0x00,0x00)
                    p2.alignment=align; cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER; set_cell_borders(cell,"CCCCCC")
            tv=total.get("var",0)
            tv_vals=["","Total",fmt_idr(total.get("mc",0)),fmt_idr(total.get("sap",0)),fmt_idr(total.get("actual",0)),fmt_idr(tv)]
            for i,(cell,val,align,w) in enumerate(zip(tbl.rows[4].cells,tv_vals,aligns,col_widths)):
                cell.width=w; p2=cell.paragraphs[0]; p2.clear(); rn=p2.add_run(val); rn.font.bold=True; rn.font.size=Pt(9)
                if i==5 and tv<0: rn.font.color.rgb=RGBColor(0xC0,0x00,0x00)
                p2.alignment=align; cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
                set_cell_bg(cell,"D6E4F0"); set_cell_borders(cell,"AAAAAA")
        perf,action=auto_remarks(sname,data,cfg)
        if perf:
            p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(3); add_run(p,perf)
        if action:
            p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(10); add_run(p,action)
        if not perf and not action:
            doc.add_paragraph().paragraph_format.space_after=Pt(10)

    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4)
    add_run(p,"C.\t",bold=True); add_run(p,"Non Moving Stock Report",bold=True)
    nm_list=[]
    if "nm_results" in st.session_state and st.session_state.nm_results:
        for nm_def in NM_DEFS:
            r=st.session_state.nm_results.get(nm_def["sloc"])
            if r: nm_list.append(r["data"])
    if not nm_list: nm_list=cfg.get("non_moving",[])
    for nm in nm_list:
        sloc_label="R&D Stores" if nm["sloc"]=="4413" else "KTC Store"
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(2)
        add_run(p,f"{sloc_label} (S.Loc {nm['sloc']})",bold=True)
        p2=doc.add_paragraph(); p2.paragraph_format.space_before=Pt(2); p2.paragraph_format.space_after=Pt(4)
        add_run(p2,f"In {quarter} {year} was recorded in total {nm['count']} material codes which in not moving "
                   f"from around 1 year to 3 years (and more) with total value IDR {fmt_idr(nm['value'])}.")
        nm_rows=nm.get("rows",[])
        if nm_rows:
            nc=[Cm(1.0),Cm(3.0),Cm(2.2),Cm(2.2),Cm(4.0),Cm(4.0),Cm(2.0)]
            na=[WD_ALIGN_PARAGRAPH.CENTER,WD_ALIGN_PARAGRAPH.LEFT,WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.CENTER,WD_ALIGN_PARAGRAPH.RIGHT,WD_ALIGN_PARAGRAPH.RIGHT,WD_ALIGN_PARAGRAPH.RIGHT]
            tbl=doc.add_table(rows=1+len(nm_rows)+1,cols=7); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl.style="Table Grid"
            nh=["No.","Criteria","MC\nNeed","MC\nDon't Need","Value Need\n(IDR)","Value Don't Need\n(IDR)","MC Total"]
            for i,(cell,text,w) in enumerate(zip(tbl.rows[0].cells,nh,nc)):
                cell.width=w; p3=cell.paragraphs[0]; p3.clear()
                r=p3.add_run(text); r.font.bold=True; r.font.size=Pt(8); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
                p3.alignment=WD_ALIGN_PARAGRAPH.CENTER; cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
                set_cell_bg(cell,"2E75B6"); set_cell_borders(cell,"FFFFFF")
            for r_i,nr in enumerate(nm_rows,start=1):
                vals=[str(r_i),nr["criteria"],str(nr["mc_need"]),str(nr["mc_dont"]),
                      fmt_idr(int(nr["val_need"])),fmt_idr(int(nr["val_dont"])),str(nr["mc_total"])]
                for i,(cell,val,align,w) in enumerate(zip(tbl.rows[r_i].cells,vals,na,nc)):
                    cell.width=w; p3=cell.paragraphs[0]; p3.clear(); rn=p3.add_run(val); rn.font.size=Pt(8)
                    p3.alignment=align; cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER; set_cell_borders(cell,"CCCCCC")
            tv=[str(sum(r["mc_need"] for r in nm_rows)),str(sum(r["mc_dont"] for r in nm_rows)),
                fmt_idr(int(sum(r["val_need"] for r in nm_rows))),fmt_idr(int(sum(r["val_dont"] for r in nm_rows))),str(nm["count"])]
            tot_v=["","Total"]+tv
            for i,(cell,val,align,w) in enumerate(zip(tbl.rows[-1].cells,tot_v,na,nc)):
                cell.width=w; p3=cell.paragraphs[0]; p3.clear(); rn=p3.add_run(val); rn.font.bold=True; rn.font.size=Pt(8)
                p3.alignment=align; cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
                set_cell_bg(cell,"D6E4F0"); set_cell_borders(cell,"AAAAAA")
            doc.add_paragraph().paragraph_format.space_after=Pt(6)

    para("Attached we also enclosed the Stock Take Details and Dead Stock Details, kindly review the report and approve.",before=10,after=4)
    para("Thank you.",before=2,after=16)
    sig_tbl=doc.add_table(rows=3,cols=3); sig_tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for col_i,(label,name) in enumerate(zip(["Prepared by,","Acknowledged by,","Approved by,"],[prepared_by,acknowledged,approved_by])):
        c0=sig_tbl.cell(0,col_i); c0.paragraphs[0].clear(); add_run(c0.paragraphs[0],label); c0.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        sig_tbl.cell(1,col_i).paragraphs[0].add_run("\n\n\n")
        c2=sig_tbl.cell(2,col_i); c2.paragraphs[0].clear(); add_run(c2.paragraphs[0],name,bold=True); c2.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    for row in sig_tbl.rows:
        for cell in row.cells: no_border(cell)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf

def generate_pdf(cfg, stores_results):
    buf=io.BytesIO()
    from reportlab.platypus import BaseDocTemplate, Frame, PageTemplate
    def _on_page(canvas,doc):
        canvas.saveState(); W,H=A4
        canvas.setFont("Helvetica-Bold",10); canvas.drawRightString(W-2.5*cm,H-1.5*cm,"Internal")
        canvas.setFont("Helvetica",9); canvas.drawCentredString(W/2,1.2*cm,f"Page {doc.page}")
        canvas.restoreState()
    doc=BaseDocTemplate(buf,pagesize=A4,leftMargin=3*cm,rightMargin=2.5*cm,topMargin=3.2*cm,bottomMargin=3.0*cm)
    frame=Frame(doc.leftMargin,doc.bottomMargin,doc.width,doc.height,id="main")
    doc.addPageTemplates([PageTemplate(id="main",frames=[frame],onPage=_on_page)])
    BLUE_MID=colors.HexColor("#2E75B6"); BLUE_LIGHT=colors.HexColor("#D6E4F0"); GREY=colors.HexColor("#F2F2F2")
    def S(name,**kw): return ParagraphStyle(name,**kw)
    s_n=S("n",fontName="Helvetica",fontSize=10,leading=16,spaceAfter=4)
    s_b=S("b",fontName="Helvetica-Bold",fontSize=10,leading=16,spaceAfter=4)
    s_bp=S("bp",fontName="Helvetica",fontSize=10,leading=15,leftIndent=0.5*cm,spaceAfter=3)
    s_sb=S("sb",fontName="Helvetica",fontSize=10,leading=15,leftIndent=1.2*cm,spaceAfter=3)
    quarter=cfg["quarter"]; year=cfg["year"]
    report_date=datetime.now().strftime("%d %B %Y").lstrip("0")
    start_date=cfg["start_date"]; end_date=cfg["end_date"]
    approved_by=cfg["approved_by"]; approved_ttl=cfg["approved_title"]
    prepared_by=cfg["prepared_by"]; acknowledged=cfg["acknowledged_by"]
    story=[]; SP=lambda h: Spacer(1,h*cm)
    story.append(Paragraph(f"Pangkalan Kerinci, {report_date}",s_n)); story.append(SP(0.3))
    story.append(Paragraph(f"<b>Subject:</b> Material Stock Take Report",s_n))
    story.append(Paragraph(f"<b>To:</b> Mr. {approved_by}",s_n))
    story.append(Paragraph(approved_ttl,s_n)); story.append(SP(0.3))
    story.append(Paragraph("Dear Sir,",s_n)); story.append(SP(0.3))
    story.append(Paragraph(f"Material Stock Take {quarter} at all R&D Stores had been done from {start_date} to {end_date}. "
        "The activity started from Analytic Lab. Store, Biomolecular Lab., GD Seed Store, "
        "KTC Store., Plant Health Lab. Store, R&D Main Store, and R&D Nursery Store.",s_n)); story.append(SP(0.4))
    story.append(Paragraph("<b>A.&nbsp;&nbsp;&nbsp;Code and Termination</b>",s_b)); story.append(SP(0.35))
    story.append(Paragraph("•  Store Location on SAP, R&D has two store locations:",s_bp))
    story.append(Paragraph("–  4413&nbsp;&nbsp;maintain all data from Genetic Deployment Seed Store, Analytic Lab. Store, Plant Health Lab. Store, Biomolecular Lab. R&D Main Store, and R&D Nursery Store.",s_sb))
    story.append(Paragraph("–  4414&nbsp;&nbsp;maintain data from RGE Tissue Culture Lab. and KTC Store.",s_sb))
    story.append(SP(0.1))
    story.append(Paragraph("•  <b>Balance</b> is where Actual Stock <i>matched</i> with SAP stock.",s_bp))
    story.append(Paragraph("•  <b>Surplus</b> is where Actual Stock <i>exceed</i> from SAP stock.",s_bp))
    story.append(Paragraph("•  <b>Deficit</b> is where Actual Stock <i>less</i> from SAP stock.",s_bp))
    story.append(SP(0.1))
    story.append(Paragraph("•  Stock Take Procedure",s_bp))
    story.append(Paragraph("–  Basis data used is from SAP.",s_sb))
    story.append(Paragraph("–  Material that has been using before Stock Take with proven manual SR will be recorded SR process.",s_sb))
    story.append(Paragraph("–  Material which failed to be shown is recorded as none.",s_sb))
    story.append(SP(0.4))
    story.append(Paragraph("<b>B.&nbsp;&nbsp;&nbsp;Stock Take Summary</b>",s_b)); story.append(SP(0.35))
    col_widths=[1.2*cm,4.8*cm,2.4*cm,3.5*cm,3.5*cm,3.0*cm]
    base_style=[("BACKGROUND",(0,0),(-1,0),BLUE_MID),("TEXTCOLOR",(0,0),(-1,0),colors.white),
                ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,-1),9),
                ("ALIGN",(0,0),(-1,0),"CENTER"),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
                ("ROWBACKGROUNDS",(0,1),(-1,-2),[colors.white,GREY]),
                ("BACKGROUND",(0,-1),(-1,-1),BLUE_LIGHT),("FONTNAME",(0,-1),(-1,-1),"Helvetica-Bold"),
                ("GRID",(0,0),(-1,-1),0.5,colors.HexColor("#CCCCCC")),
                ("ALIGN",(2,1),(-1,-1),"RIGHT"),("ALIGN",(0,1),(0,-1),"CENTER"),
                ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
                ("LEFTPADDING",(0,0),(-1,-1),5),("RIGHTPADDING",(0,0),(-1,-1),5)]
    for idx,store_def in enumerate(STORE_DEFS,start=1):
        sname=store_def["name"]; result=stores_results.get(store_def["key"])
        if result is None: continue
        data=result["data"]; overrides=cfg.get("store_overrides",{}).get(sname,{})
        pic=overrides.get("pic",store_def["default_pic"])
        date=overrides.get("stock_take_date") or data.get("date_from_file") or cfg["start_date"]
        story.append(SP(0.25))
        story.append(Paragraph(f"<b>{idx}.&nbsp;&nbsp;{sname}, {date}, PIC {pic}.</b>",s_n)); story.append(SP(0.2))
        story.append(Paragraph(build_intro(sname,data,store_def,cfg),s_n)); story.append(SP(0.25))
        total=data.get("total",{}); is_seed=store_def["seed"]
        if is_seed:
            sr=data.get("species_rows",[])
            sw=[0.55*cm,1.3*cm,3.6*cm,0.9*cm,1.7*cm,1.7*cm,1.2*cm,2.5*cm,2.5*cm,1.7*cm,1.3*cm]
            sh=["No.","MC","Material Description","Unit","SAP Qty","Total Actual\nQty. (g)","Var.\n(g)","SAP Value\n(IDR)","Actual Value\n(IDR)","Var. (IDR)","Code"]
            td=[sh]
            for s in sr:
                vq=s["var_qty"]; vv=s["var"]
                td.append([str(s["no"]),str(s["mc"]),s["desc"],s["unit"],fmt_idr(s["sap_qty"]),fmt_idr(s["actual_qty"]),
                           "-" if vq==0 else fmt_idr(vq),fmt_idr(s["sap"]),fmt_idr(s["actual"]),"-" if vv==0 else fmt_idr(vv),s["code"].capitalize()])
            tot=total; tvq=tot.get("var_qty",0); tv=tot.get("var",0)
            td.append(["","Total","","",fmt_idr(tot.get("sap_qty",0)),fmt_idr(tot.get("actual_qty",0)),
                       "-" if tvq==0 else fmt_idr(tvq),fmt_idr(tot.get("sap",0)),fmt_idr(tot.get("actual",0)),
                       "-" if tv==0 else fmt_idr(tv),""])
            ss=[("BACKGROUND",(0,0),(-1,0),BLUE_MID),("TEXTCOLOR",(0,0),(-1,0),colors.white),
                ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,-1),7),
                ("ALIGN",(0,0),(-1,0),"CENTER"),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
                ("ROWBACKGROUNDS",(0,1),(-1,-2),[colors.white,GREY]),
                ("BACKGROUND",(0,-1),(-1,-1),BLUE_LIGHT),("FONTNAME",(0,-1),(-1,-1),"Helvetica-Bold"),
                ("GRID",(0,0),(-1,-1),0.5,colors.HexColor("#CCCCCC")),
                ("ALIGN",(4,1),(-1,-1),"RIGHT"),("ALIGN",(0,1),(3,-1),"CENTER"),
                ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
                ("LEFTPADDING",(0,0),(-1,-1),3),("RIGHTPADDING",(0,0),(-1,-1),3)]
            if tv<0: ss.append(("TEXTCOLOR",(9,-1),(9,-1),colors.red))
            for r_i,s in enumerate(sr,start=1):
                if s["var"]<0: ss.append(("TEXTCOLOR",(9,r_i),(9,r_i),colors.red))
            tbl=Table(td,colWidths=sw,repeatRows=1); tbl.setStyle(TableStyle(ss))
            story.append(tbl); story.append(SP(0.25))
        else:
            rs=[("Balance Items",data.get("balance",{})),("Surplus Items",data.get("surplus",{})),("Deficit Items",data.get("deficit",{}))]
            td=[["No.","Description","Number\nof MC","SAP Value\n(IDR)","Actual Value\n(IDR)","Var. (IDR)"]]
            for r_i,(desc,rd) in enumerate(rs,start=1):
                vv=rd.get("var",0); mc_d=fmt_idr(rd.get("mc",0)) if rd.get("mc",0)!=0 else "-"
                vd=f"({fmt_idr(abs(vv))})" if vv<0 else fmt_idr(vv)
                td.append([str(r_i),desc,mc_d,fmt_idr(rd.get("sap",0)),fmt_idr(rd.get("actual",0)),vd])
            tv=total.get("var",0); tvd=f"({fmt_idr(abs(tv))})" if tv<0 else fmt_idr(tv)
            td.append(["","Total",fmt_idr(total.get("mc",0)),fmt_idr(total.get("sap",0)),fmt_idr(total.get("actual",0)),tvd])
            sc=list(base_style)
            for r_i,(_,rd) in enumerate(rs,start=1):
                if rd.get("var",0)<0: sc.append(("TEXTCOLOR",(5,r_i),(5,r_i),colors.red))
            if tv<0: sc.append(("TEXTCOLOR",(5,-1),(5,-1),colors.red))
            tbl=Table(td,colWidths=col_widths,repeatRows=1); tbl.setStyle(TableStyle(sc))
            story.append(tbl); story.append(SP(0.35))
        perf,action=auto_remarks(sname,data,cfg)
        if perf:   story.append(Paragraph(perf,s_n)); story.append(SP(0.1))
        if action: story.append(Paragraph(action,s_n))
        story.append(SP(0.45))
    story.append(Paragraph("<b>C.&nbsp;&nbsp;&nbsp;Non Moving Stock Report</b>",s_b)); story.append(SP(0.35))
    nm_list_pdf=[]
    if "nm_results" in st.session_state and st.session_state.nm_results:
        for nm_def in NM_DEFS:
            r=st.session_state.nm_results.get(nm_def["sloc"])
            if r: nm_list_pdf.append(r["data"])
    if not nm_list_pdf: nm_list_pdf=cfg.get("non_moving",[])
    nm_col_w_pdf=[0.8*cm,2.5*cm,1.8*cm,1.8*cm,3.5*cm,3.5*cm,1.8*cm]
    nm_base_style=[("BACKGROUND",(0,0),(-1,0),BLUE_MID),("TEXTCOLOR",(0,0),(-1,0),colors.white),
                   ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,-1),8),
                   ("ALIGN",(0,0),(-1,0),"CENTER"),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
                   ("ROWBACKGROUNDS",(0,1),(-1,-2),[colors.white,colors.HexColor("#F2F2F2")]),
                   ("BACKGROUND",(0,-1),(-1,-1),BLUE_LIGHT),("FONTNAME",(0,-1),(-1,-1),"Helvetica-Bold"),
                   ("GRID",(0,0),(-1,-1),0.5,colors.HexColor("#CCCCCC")),
                   ("ALIGN",(2,1),(-1,-1),"RIGHT"),("ALIGN",(0,1),(1,-1),"CENTER"),
                   ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
                   ("LEFTPADDING",(0,0),(-1,-1),4),("RIGHTPADDING",(0,0),(-1,-1),4)]
    for nm in nm_list_pdf:
        sloc_label="R&D Stores" if nm["sloc"]=="4413" else "KTC Store"
        story.append(Paragraph(f"<b>{sloc_label} (S.Loc {nm['sloc']})</b>",s_b)); story.append(SP(0.15))
        story.append(Paragraph(f"In {quarter} {year} was recorded in total {nm['count']} material codes which in not moving "
                               f"from around 1 year to 3 years (and more) with total value IDR {fmt_idr(nm['value'])}.",s_n))
        story.append(SP(0.2))
        nm_rows=nm.get("rows",[])
        if nm_rows:
            td=[["No.","Criteria","MC\nNeed","MC\nDon't Need","Value Need\n(IDR)","Value Don't Need\n(IDR)","MC Total"]]
            for r_i,nr in enumerate(nm_rows,start=1):
                td.append([str(r_i),nr["criteria"],str(nr["mc_need"]),str(nr["mc_dont"]),
                           fmt_idr(int(nr["val_need"])),fmt_idr(int(nr["val_dont"])),str(nr["mc_total"])])
            td.append(["","Total",str(sum(r["mc_need"] for r in nm_rows)),str(sum(r["mc_dont"] for r in nm_rows)),
                       fmt_idr(int(sum(r["val_need"] for r in nm_rows))),
                       fmt_idr(int(sum(r["val_dont"] for r in nm_rows))),str(nm["count"])])
            tbl=Table(td,colWidths=nm_col_w_pdf,repeatRows=1); tbl.setStyle(TableStyle(nm_base_style))
            story.append(tbl)
        story.append(SP(0.35))
    story.append(SP(0.4))
    story.append(Paragraph("Attached we also enclosed the Stock Take Details and Dead Stock Details, kindly review the report and approve.",s_n))
    story.append(SP(0.2)); story.append(Paragraph("Thank you.",s_n)); story.append(SP(0.8))
    sig_data=[["Prepared by,","Acknowledged by,","Approved by,"],["","",""],["","",""],["","",""],
              [prepared_by,acknowledged,approved_by]]
    sig_tbl=Table(sig_data,colWidths=[5.5*cm,5.5*cm,5.5*cm])
    sig_tbl.setStyle(TableStyle([("ALIGN",(0,0),(-1,-1),"CENTER"),("FONTNAME",(0,-1),(-1,-1),"Helvetica-Bold"),
                                  ("FONTSIZE",(0,0),(-1,-1),10),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
                                  ("LINEABOVE",(0,-1),(-1,-1),0.5,colors.black)]))
    story.append(sig_tbl); doc.build(story); buf.seek(0); return buf

# ══════════════════════════════════════════════════════════════════════════════
# UI
# ══════════════════════════════════════════════════════════════════════════════
if "stores_results" not in st.session_state: st.session_state.stores_results={}
if "nm_results"     not in st.session_state: st.session_state.nm_results={}
if "detected_q"     not in st.session_state: st.session_state.detected_q=None
if "detected_y"     not in st.session_state: st.session_state.detected_y=None

cfg=load_config()
if st.session_state.detected_q: cfg["quarter"]=st.session_state.detected_q
if st.session_state.detected_y: cfg["year"]=st.session_state.detected_y

tab_upload,tab_config,tab_generate=st.tabs(["📁 Upload Store Files","⚙️ Configuration","📄 Generate Report"])

with tab_upload:
    st.markdown("### Upload Store Excel Files")
    st.markdown("""<div class="warn">Upload all 9 Excel files (7 store + 2 Non-Moving Stock). Each must have a sheet named <b>auto</b>.<br>
Store: <code>PHL</code> · <code>SOL</code> · <code>BIO</code> · <code>RDS</code> · <code>KRN</code> · <code>SEED</code> · <code>TCL</code> &nbsp;|&nbsp;
Non-Moving: filename must contain <code>Non Moving</code> + <code>4413</code> or <code>4414</code></div>""", unsafe_allow_html=True)

    uploaded_files=st.file_uploader("Upload Excel files",type=["xlsx","xls"],accept_multiple_files=True)
    if uploaded_files:
        for f in uploaded_files:
            nm_def=detect_nm_file(f.name)
            if nm_def:
                data,err=parse_nm_file(f,nm_def)
                if err:
                    st.error(f"❌ **{f.name}** (Non-Moving {nm_def['sloc']}) — {err}")
                    st.session_state.nm_results.pop(nm_def["sloc"],None)
                else:
                    st.session_state.nm_results[nm_def["sloc"]]={"data":data,"filename":f.name}
                    q=extract_quarter_from_filename(f.name); y=extract_year_from_filename(f.name)
                    if q: st.session_state.detected_q=q; cfg["quarter"]=q
                    if y: st.session_state.detected_y=y; cfg["year"]=y
                continue
            store_def=detect_store(f.name)
            if store_def is None:
                st.error(f"❌ **{f.name}** — cannot detect (expected: PHL/SOL/BIO/RDS/KRN/SEED/TCL or Non-Moving with 4413/4414)")
                continue
            data,err=parse_store_file(f,store_def)
            if err:
                st.error(f"❌ **{f.name}** ({store_def['name']}) — {err}")
                st.session_state.stores_results.pop(store_def["key"],None)
            else:
                st.session_state.stores_results[store_def["key"]]={"data":data,"filename":f.name,"store_def":store_def}
                q=extract_quarter_from_filename(f.name); y=extract_year_from_filename(f.name)
                if q: st.session_state.detected_q=q; cfg["quarter"]=q
                if y: st.session_state.detected_y=y; cfg["year"]=y

    # ── Auto-detect date range from all uploaded store filenames ──────────────
    all_dates = []
    for key, result in st.session_state.stores_results.items():
        d = result["data"].get("date_from_file")
        if d:
            try:
                from datetime import datetime as dt
                parsed = dt.strptime(d, "%d-%b-%Y")
                all_dates.append((parsed, d))
            except: pass
    if all_dates:
        all_dates.sort(key=lambda x: x[0])
        earliest = all_dates[0][1]
        latest   = all_dates[-1][1]
        if cfg.get("start_date") != earliest or cfg.get("end_date") != latest:
            cfg["start_date"] = earliest
            cfg["end_date"]   = latest
            save_config(cfg)

    if st.session_state.detected_q or st.session_state.detected_y:
        st.info(f"📅 Auto-detected — Quarter: **{cfg.get('quarter','?')}** | Year: **{cfg.get('year','?')}**")
    if all_dates and len(set(d for _,d in all_dates)) > 1:
        st.info(f"📅 Stock take period auto-detected — **{cfg['start_date']}** to **{cfg['end_date']}**")
    elif all_dates:
        st.info(f"📅 Stock take date auto-detected — **{cfg['start_date']}**")

    st.markdown("---")
    store_loaded=[sd for sd in STORE_DEFS if sd["key"] in st.session_state.stores_results]
    nm_loaded_list=[nm for nm in NM_DEFS if nm["sloc"] in st.session_state.nm_results]
    total_ok=len(store_loaded)+len(nm_loaded_list)
    total_all=len(STORE_DEFS)+len(NM_DEFS)
    if total_ok==total_all:
        st.markdown(f'<div class="ok">✅ All {total_all} files loaded and ready!</div>',unsafe_allow_html=True)
    else:
        st.markdown(f'<div class="warn">📂 {total_ok}/{total_all} files loaded.</div>',unsafe_allow_html=True)

    st.markdown("##### 🏪 Store Files (7)")
    for sd in STORE_DEFS:
        result=st.session_state.stores_results.get(sd["key"])
        c1,c2,c3=st.columns([3,2,5]); c1.markdown(f"**{sd['name']}**")
        if result:
            dr=result["data"]; tot=dr.get("total",{}); var=tot.get("var",0)
            c2.markdown('<span class="store-ok">✅ Loaded</span>',unsafe_allow_html=True)
            vt="✅ Zero Var" if var==0 else (f"🔼 Surplus: {fmt_idr(var)}" if var>0 else f"🔽 Deficit: {fmt_idr(abs(var))}")
            date_detected=dr.get("date_from_file","")
            date_txt=f" &nbsp;|&nbsp; 📅 {date_detected}" if date_detected else ""
            c3.markdown(f"MC: **{tot.get('mc',0)}** &nbsp;|&nbsp; SAP: {fmt_idr(tot.get('sap',0))} &nbsp;|&nbsp; {vt}{date_txt}")
        else:
            c2.markdown('<span class="store-err">⏳ Missing</span>',unsafe_allow_html=True)
            c3.markdown(f"<small>Upload file with keyword <code>{sd['key']}</code></small>",unsafe_allow_html=True)

    st.markdown("##### 📦 Non-Moving Stock Files (2)")
    for nm_def in NM_DEFS:
        result=st.session_state.nm_results.get(nm_def["sloc"])
        c1,c2,c3=st.columns([3,2,5]); c1.markdown(f"**S.Loc {nm_def['sloc']}** — {nm_def['desc']}")
        if result:
            d=result["data"]; c2.markdown('<span class="store-ok">✅ Loaded</span>',unsafe_allow_html=True)
            c3.markdown(f"MC: **{d['count']}** &nbsp;|&nbsp; Value: IDR {fmt_idr(d['value'])} &nbsp;|&nbsp; <small><i>{result['filename']}</i></small>",unsafe_allow_html=True)
        else:
            c2.markdown('<span class="store-err">⏳ Missing</span>',unsafe_allow_html=True)
            c3.markdown(f"<small>Upload <code>Non Moving Stock Report {nm_def['sloc']} - Q?.xlsx</code></small>",unsafe_allow_html=True)

with tab_config:
    st.markdown("### Report Configuration")
    changed=False
    if st.session_state.detected_q or st.session_state.detected_y:
        st.info(f"📅 Auto-detected from filenames — Quarter: **{cfg.get('quarter','?')}** | Year: **{cfg.get('year','?')}**")

    st.markdown("#### 📅 Stock Take Period")
    st.caption("Auto-filled from filenames (earliest → latest date). Override manually if needed.")
    c1,c2=st.columns(2)
    with c1:
        v=st.text_input("Stock Take Start Date",cfg["start_date"],help="e.g. 22-Sep-2025")
        if v!=cfg["start_date"]: cfg["start_date"]=v; changed=True
    with c2:
        v=st.text_input("Stock Take End Date",cfg["end_date"])
        if v!=cfg["end_date"]: cfg["end_date"]=v; changed=True

    st.markdown("#### 👤 Signatories")
    c1,c2,c3=st.columns(3)
    with c1:
        v=st.text_input("Prepared By",cfg["prepared_by"])
        if v!=cfg["prepared_by"]: cfg["prepared_by"]=v; changed=True
    with c2:
        v=st.text_input("Acknowledged By",cfg["acknowledged_by"])
        if v!=cfg["acknowledged_by"]: cfg["acknowledged_by"]=v; changed=True
    with c3:
        v=st.text_input("Approved By",cfg["approved_by"])
        if v!=cfg["approved_by"]: cfg["approved_by"]=v; changed=True
        v=st.text_input("Approved By Title",cfg["approved_title"])
        if v!=cfg["approved_title"]: cfg["approved_title"]=v; changed=True

    st.markdown("#### 🏪 Per-Store Overrides")
    if "store_overrides" not in cfg: cfg["store_overrides"]={}
    for sd in STORE_DEFS:
        sname=sd["name"]
        if sname not in cfg["store_overrides"]: cfg["store_overrides"][sname]={}
        ov=cfg["store_overrides"][sname]
        with st.expander(f"**{sname}**",expanded=False):
            c1,c2=st.columns(2)
            with c1:
                v=st.text_input("PIC",ov.get("pic",sd["default_pic"]),key=f"pic_{sd['key']}")
                if v!=ov.get("pic",sd["default_pic"]): ov["pic"]=v; changed=True
            with c2:
                v=st.text_input("Stock Take Date",ov.get("stock_take_date",cfg["start_date"]),key=f"date_{sd['key']}")
                if v!=ov.get("stock_take_date",cfg["start_date"]): ov["stock_take_date"]=v; changed=True
            st.caption("✏️ Custom remarks (leave blank for auto-generated)")
            v=st.text_area("Performance Note",ov.get("performance_note",""),height=60,key=f"perf_{sd['key']}")
            if v!=ov.get("performance_note",""): ov["performance_note"]=v if v else None; changed=True
            v=st.text_area("Action Note",ov.get("action_note",""),height=60,key=f"act_{sd['key']}")
            if v!=ov.get("action_note",""): ov["action_note"]=v if v else None; changed=True
            if sd["seed"]:
                result=st.session_state.stores_results.get(sd["key"])
                if result:
                    d=result["data"]
                    st.info(f"🌱 Seed Lots: **{d.get('seed_lots_count','?')}** | Species: **{d.get('species_count','?')}** | Weight: **{d.get('total_weight_tons','?')}**")
                else:
                    st.caption("🌱 Upload SEED file to auto-compute.")

    st.markdown("#### 📦 Non-Moving Stock (fallback if files not uploaded)")
    nm=cfg.get("non_moving",[{"sloc":"4413","desc":"R&D Stores","count":0,"value":0},
                               {"sloc":"4414","desc":"KTC Store","count":0,"value":0}])
    for i,entry in enumerate(nm):
        c1,c2,c3,c4=st.columns([2,3,2,3])
        with c1:
            v=st.text_input("S.Loc",entry["sloc"],key=f"nm_sloc_{i}")
            if v!=entry["sloc"]: nm[i]["sloc"]=v; changed=True
        with c2:
            v=st.text_input("Description",entry["desc"],key=f"nm_desc_{i}")
            if v!=entry["desc"]: nm[i]["desc"]=v; changed=True
        with c3:
            v=st.number_input("MC Count",value=entry["count"],key=f"nm_cnt_{i}",min_value=0)
            if v!=entry["count"]: nm[i]["count"]=v; changed=True
        with c4:
            v=st.number_input("Total Value (IDR)",value=entry["value"],key=f"nm_val_{i}",min_value=0,step=1000000)
            if v!=entry["value"]: nm[i]["value"]=v; changed=True
    cfg["non_moving"]=nm
    if st.button("💾 Save Configuration",type="primary"):
        save_config(cfg); st.success("✅ Saved!")
    elif changed:
        save_config(cfg)

with tab_generate:
    st.markdown("### Generate Report")
    loaded=[sd for sd in STORE_DEFS if sd["key"] in st.session_state.stores_results]
    missing=[sd for sd in STORE_DEFS if sd["key"] not in st.session_state.stores_results]
    nm_missing=[nm for nm in NM_DEFS if nm["sloc"] not in st.session_state.nm_results]
    if missing:
        st.markdown(f'<div class="warn">⚠️ Missing: {", ".join(sd["name"] for sd in missing)}</div>',unsafe_allow_html=True)
    if nm_missing:
        st.markdown(f'<div class="warn">⚠️ Missing NM files: {", ".join("S.Loc "+nm["sloc"] for nm in nm_missing)} — config fallback used.</div>',unsafe_allow_html=True)
    if loaded:
        q=cfg.get("quarter","?"); y=cfg.get("year","?")
        st.markdown(f'<div class="ok">✅ {len(loaded)}/7 stores — Quarter: <b>{q}</b> | Year: <b>{y}</b></div>',unsafe_allow_html=True)
    if st.session_state.stores_results:
        with st.expander("👁️ Preview auto-generated remarks"):
            for sd in STORE_DEFS:
                r=st.session_state.stores_results.get(sd["key"])
                if not r: continue
                perf,action=auto_remarks(sd["name"],r["data"],cfg)
                st.markdown(f"**{sd['name']}**")
                if perf:   st.info(f"📝 {perf}")
                if action: st.warning(f"⚠️ {action}")
                if not perf and not action: st.caption("_(blank — admin fills after download)_")
    q=cfg.get("quarter","Q?"); y=cfg.get("year","????")
    fname=f"Stock_Take_Report_{q}_{y}"
    c1,c2=st.columns(2)
    with c1:
        if st.button("📝 Generate Word (.docx)",use_container_width=True,type="primary",disabled=not loaded):
            with st.spinner("Generating..."):
                buf=generate_docx(cfg,st.session_state.stores_results)
                st.download_button(f"⬇️ Download {fname}.docx",buf,file_name=f"{fname}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",use_container_width=True)
    with c2:
        if st.button("📄 Generate PDF",use_container_width=True,type="primary",disabled=not loaded):
            with st.spinner("Generating..."):
                buf=generate_pdf(cfg,st.session_state.stores_results)
                st.download_button(f"⬇️ Download {fname}.pdf",buf,file_name=f"{fname}.pdf",
                    mime="application/pdf",use_container_width=True)